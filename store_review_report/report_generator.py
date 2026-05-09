#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
门店经营复盘报告生成器 v2.0
===========================
数据驱动 · 排版精美 · 逻辑严谨 · 通用复用

使用方法：
    1. 按数据格式准备门店数据（参考示例 data_xiangdong）
    2. 调用 StoreReviewReport(data).generate(output_path)
    3. 自动生成精美的 Word 复盘报告

设计说明：
    - 配色：深靛蓝(主色) + 石板灰(辅助) + 翠绿/琥珀/砖红(状态色)
    - 逻辑：结论先行 → 数据支撑 → 问题拆解 → 改善行动 → 下月计划
    - 通用：所有门店数据均可按同一结构填入，自动计算达成率、环比、利润等
"""

from dataclasses import dataclass, field
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ==================== 配色方案 ====================
COLORS = {
    "primary":      "1B2A4A",   # 深靛蓝 - 主标题、表头
    "secondary":    "334155",   # 石板灰 - 副标题
    "accent":       "B45309",   # 琥珀金 - 装饰、强调
    "success":      "047857",   # 翠绿 - 达成优秀
    "warning":      "B45309",   # 琥珀 - 达成一般
    "danger":       "BE123C",   # 砖红 - 达成差、亏损
    "bg_light":     "F8FAFC",   # 极浅灰 - 表格交替背景
    "bg_accent":    "EFF6FF",   # 浅蓝灰 - 强调区块背景
    "bg_warm":      "FFFBEB",   # 暖浅黄 - 提示框背景
    "text":         "1E293B",   # 正文深灰
    "text_muted":   "64748B",   # 次要文字
    "white":        "FFFFFF",
    "border":       "CBD5E1",   # 边框浅灰
}


# ==================== 数据模型 ====================
@dataclass
class CategoryData:
    """品类数据"""
    name: str           # 品类名称
    target: float       # 本月目标
    actual: float       # 实际达成
    last_month: float   # 上月实绩
    unit: str = "元"    # 单位（元/台/%/人）
    is_highlight: bool = False  # 是否高亮（核心指标）


@dataclass
class CostItem:
    """费用项目"""
    category: str       # 大类：固定费用/人力费用/资金费用/运营费用/变动费用
    name: str           # 明细名称
    amount: float       # 金额（元）
    review: str = ""    # 费用审视说明


@dataclass
class OperationMetric:
    """运营指标"""
    name: str
    last_month: str
    current_month: str
    diagnosis: str      # 诊断说明


@dataclass
class ActionItem:
    """改善动作"""
    category: str       # 分类：手机板块/短板品类/优势守住/过程管理
    action: str         # 动作描述
    standard: str       # 量化标准
    owner: str          # 责任人
    deadline: str       # 完成期限


@dataclass
class StoreReviewData:
    """门店复盘数据总包"""
    store_name: str
    review_month: str       # 如：2026年4月
    author: str
    date: str               # 编制日期

    # 核心指标
    gmv: float
    gmv_target: float
    gmv_last_month: float
    profit: float           # 净利润（可为负）
    cost_total: float
    rebate_amount: float    # 返利金额
    rebate_rate: float      # 返利率（如 0.1008）

    # 明细数据
    categories: List[CategoryData]
    costs: List[CostItem]
    operation_metrics: List[OperationMetric]

    # 分析结论
    conclusions: List[str]          # 经营摘要结论
    highlights: List[str]           # 亮点说明
    weaknesses: List[str]           # 短板说明
    problems: List[str]             # 问题拆解要点
    five_whys: List[tuple]          # [(问题, 回答), ...]

    # 行动计划
    actions: List[ActionItem]

    # 下月计划
    next_month_goals: List[dict]    # [{"维度": str, "基础目标": str, "挑战目标": str, "策略": str}]

    # 经营者日记
    reflection: str


# ==================== 工具函数 ====================
def hex_to_rgb(hex_str: str) -> RGBColor:
    """十六进制颜色转 RGBColor"""
    hex_str = hex_str.lstrip('#')
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def set_cell_shading(cell, color_hex: str):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex.lstrip('#'))
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, color: str = "CBD5E1", size: str = "4", style: str = "single"):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), style)
        edge_el.set(qn('w:sz'), size)
        edge_el.set(qn('w:color'), color.lstrip('#'))
        tcBorders.append(edge_el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, left: int = 100, right: int = 100):
    """设置单元格内边距（单位：dxa，1/20磅）"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin = OxmlElement(f'w:{edge}')
        margin.set(qn('w:w'), str(val))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)


def set_run_font(run, name: str = "Microsoft YaHei", size: int = 10, bold: bool = False,
                 color: str = "1E293B", italic: bool = False):
    """统一设置 run 字体"""
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = hex_to_rgb(color)


def add_horizontal_line(doc, color: str = "CBD5E1", width: int = 60):
    """添加装饰水平线（模拟）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * width)
    run.font.color.rgb = hex_to_rgb(color)
    run.font.size = Pt(8)


# ==================== 报告生成器 ====================
class StoreReviewReport:
    def __init__(self, data: StoreReviewData):
        self.data = data
        self.doc = Document()
        self._setup_document()

    def _setup_document(self):
        """文档基础设置：字体、页边距、页脚"""
        style = self.doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        style.font.size = Pt(10.5)
        style.font.color.rgb = hex_to_rgb(COLORS["text"])

        # 页边距：窄边距，更多内容空间
        for section in self.doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)

        # 添加页脚（页码）
        self._add_page_footer()

    def _add_page_footer(self):
        """添加页脚：页码 + 门店信息"""
        for section in self.doc.sections:
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.paragraph_format.space_before = Pt(6)

            # 清空现有内容
            footer_para.clear()

            # 添加细线分隔
            run_line = footer_para.add_run("─" * 50)
            set_run_font(run_line, size=7, color=COLORS["border"])

            # 新段落放页码信息
            footer_para2 = footer.add_paragraph()
            footer_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para2.paragraph_format.space_after = Pt(4)
            run_info = footer_para2.add_run(
                f"{self.data.store_name} · {self.data.review_month} 经营复盘报告"
            )
            set_run_font(run_info, size=8, color=COLORS["text_muted"])

    def _add_heading(self, text: str, level: int = 1, color: str = None, align=WD_ALIGN_PARAGRAPH.LEFT):
        """添加层级标题，带装饰线效果"""
        color = color or COLORS["primary"]
        sizes = {1: 16, 2: 13, 3: 11}
        size = sizes.get(level, 11)

        p = self.doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.3

        # 左侧装饰竖条（用字符模拟）
        if level == 1:
            run_bar = p.add_run("▌ ")
            set_run_font(run_bar, size=size + 2, bold=True, color=color)

        run = p.add_run(text)
        set_run_font(run, size=size, bold=True, color=color)

        # 标题下方细线（用下边框段落模拟）
        if level == 1:
            p_bottom = self.doc.add_paragraph()
            p_bottom.paragraph_format.space_before = Pt(0)
            p_bottom.paragraph_format.space_after = Pt(8)
            run_line = p_bottom.add_run("")
            # 这里通过段落下边框实现，但 python-docx 原生不支持段落边框，暂时用空段落留白

    def _add_para(self, text: str, bold: bool = False, italic: bool = False, size: int = 10.5,
                  color: str = None, align=WD_ALIGN_PARAGRAPH.LEFT, indent: bool = True,
                  space_after: int = 4):
        """添加标准段落"""
        color = color or COLORS["text"]
        p = self.doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.line_spacing = 1.6
        p.paragraph_format.space_after = Pt(space_after)
        if indent and align == WD_ALIGN_PARAGRAPH.LEFT:
            p.paragraph_format.first_line_indent = Cm(0.74)

        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)
        return p

    def _add_bullet(self, text: str, bold: bool = False, color: str = None, icon: str = "•"):
        """添加项目符号"""
        color = color or COLORS["text"]
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.6
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.4)

        run_icon = p.add_run(f"{icon}  ")
        set_run_font(run_icon, size=10, bold=True, color=color)

        run_text = p.add_run(text)
        set_run_font(run_text, size=10, bold=bold, color=COLORS["text"])

    def _make_table_header(self, table, headers: List[str], color: str = None):
        """生成统一表头"""
        color = color or COLORS["primary"]
        row = table.rows[0]
        for i, h in enumerate(headers):
            cell = row.cells[i]
            set_cell_shading(cell, color)
            set_cell_border(cell, color=color, size="6")
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            set_run_font(run, size=10.5, bold=True, color=COLORS["white"])

    def _make_table_row(self, table, row_idx: int, values: List[str],
                        aligns: List[WD_ALIGN_PARAGRAPH] = None,
                        shadings: List[str] = None,
                        bolds: List[bool] = None):
        """生成统一数据行"""
        aligns = aligns or [WD_ALIGN_PARAGRAPH.CENTER] * len(values)
        shadings = shadings or [None] * len(values)
        bolds = bolds or [False] * len(values)

        row = table.rows[row_idx]
        for j, text in enumerate(values):
            cell = row.cells[j]
            set_cell_border(cell, color=COLORS["border"], size="4")
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            if shadings[j]:
                set_cell_shading(cell, shadings[j])

            p = cell.paragraphs[0]
            p.alignment = aligns[j]
            run = p.add_run(str(text))
            set_run_font(run, size=10, bold=bolds[j], color=COLORS["text"])

    def _rate_color(self, rate: float) -> str:
        """根据达成率返回颜色"""
        if rate >= 1.0:
            return COLORS["success"]
        elif rate >= 0.75:
            return COLORS["warning"]
        else:
            return COLORS["danger"]

    def _rate_bg(self, rate: float) -> str:
        """根据达成率返回背景色"""
        if rate >= 1.0:
            return "D1FAE5"  # 浅绿
        elif rate >= 0.75:
            return "FEF3C7"  # 浅黄
        else:
            return "FFE4E6"  # 浅红

    # ==================== 封面 ====================
    def add_cover(self):
        """精美封面 - 采用居中对齐、层次分明的设计"""
        for _ in range(4):
            self.doc.add_paragraph()

        # 顶部装饰细线
        p_top = self.doc.add_paragraph()
        p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_top.paragraph_format.space_after = Pt(20)
        run_top = p_top.add_run("━━━━━━━━━━━━━━━━━━━━")
        set_run_font(run_top, size=12, color=COLORS["accent"])

        # 主标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run("门店月度经营复盘报告")
        set_run_font(run, size=28, bold=True, color=COLORS["primary"])

        # 英文副标题（增加设计感）
        p_en = self.doc.add_paragraph()
        p_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_en.paragraph_format.space_after = Pt(6)
        run_en = p_en.add_run("MONTHLY BUSINESS REVIEW")
        set_run_font(run_en, size=10, color=COLORS["text_muted"])
        run_en.font.name = "Arial"
        run_en._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial")

        # 门店月份标签
        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(20)
        run2 = p2.add_run(f"{self.data.store_name}  ·  {self.data.review_month}")
        set_run_font(run2, size=14, color=COLORS["secondary"])

        # 中部装饰线
        p_mid = self.doc.add_paragraph()
        p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_mid.paragraph_format.space_after = Pt(30)
        run_mid = p_mid.add_run("◆  ◆  ◆")
        set_run_font(run_mid, size=10, color=COLORS["accent"])

        for _ in range(2):
            self.doc.add_paragraph()

        # 信息卡片（表格实现，标签深蓝、值浅灰，精致边框）
        info_table = self.doc.add_table(rows=4, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_data = [
            ("门店名称", self.data.store_name),
            ("复盘周期", self.data.review_month),
            ("编制人", self.data.author),
            ("编制日期", self.data.date),
        ]
        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            cell0 = row.cells[0]
            cell1 = row.cells[1]

            # 标签列：深蓝背景白字
            set_cell_shading(cell0, COLORS["primary"])
            set_cell_border(cell0, color=COLORS["primary"], size="8")
            set_cell_margins(cell0, top=140, bottom=140, left=180, right=180)
            p0 = cell0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r0 = p0.add_run(label)
            set_run_font(r0, size=11, bold=True, color=COLORS["white"])

            # 值列：极浅灰背景
            set_cell_shading(cell1, "F1F5F9")
            set_cell_border(cell1, color=COLORS["primary"], size="8")
            set_cell_margins(cell1, top=140, bottom=140, left=180, right=180)
            p1 = cell1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p1.add_run(value)
            set_run_font(r1, size=11, bold=True, color=COLORS["primary"])

        for row in info_table.rows:
            row.cells[0].width = Cm(3.8)
            row.cells[1].width = Cm(6.2)

        # 底部留白
        for _ in range(3):
            self.doc.add_paragraph()

        # 底部装饰文字
        p_bottom = self.doc.add_paragraph()
        p_bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_bottom = p_bottom.add_run("数据驱动经营 · 复盘指引未来")
        set_run_font(run_bottom, size=10, italic=True, color=COLORS["text_muted"])

        self.doc.add_page_break()

    # ==================== 第一章：经营摘要 ====================
    def add_summary(self):
        self._add_heading("一、经营摘要", level=1)
        self._add_para("本月经营的核心结论如下，后续章节将用数据逐一支撑：",
                        italic=True, color=COLORS["text_muted"], indent=False)

        for i, conclusion in enumerate(self.data.conclusions, 1):
            self._add_bullet(conclusion, bold=(i <= 2), color=COLORS["primary"], icon=f"{i}.")

        # 核心指标速览卡片（用表格实现）
        self.doc.add_paragraph()
        quick_table = self.doc.add_table(rows=1, cols=4)
        quick_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["销售额(GMV)", "净利润", "费用总额", "返利率"]
        values = [
            f"¥{self.data.gmv:,.0f}",
            f"¥{self.data.profit:,.0f}",
            f"¥{self.data.cost_total:,.0f}",
            f"{self.data.rebate_rate*100:.2f}%"
        ]
        # 利润颜色
        profit_color = COLORS["success"] if self.data.profit >= 0 else COLORS["danger"]
        profit_bg = "D1FAE5" if self.data.profit >= 0 else "FFE4E6"
        bgs = [COLORS["bg_light"], profit_bg, COLORS["bg_light"], COLORS["bg_light"]]

        for i, (h, v) in enumerate(zip(headers, values)):
            cell = quick_table.rows[0].cells[i]
            set_cell_shading(cell, bgs[i])
            set_cell_border(cell, color=COLORS["border"], size="6")
            set_cell_margins(cell, top=140, bottom=140, left=120, right=120)

            p_h = cell.paragraphs[0]
            p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_h = p_h.add_run(h)
            set_run_font(r_h, size=9, color=COLORS["text_muted"])

            p_v = cell.add_paragraph()
            p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_v = p_v.add_run(v)
            c = profit_color if i == 1 else COLORS["primary"]
            set_run_font(r_v, size=14, bold=True, color=c)

        for cell in quick_table.rows[0].cells:
            cell.width = Cm(3.5)

        self.doc.add_paragraph()

    # ==================== 第二章：核心经营数据 ====================
    def add_sales_analysis(self):
        self._add_heading("二、核心经营数据", level=1)
        self._add_para("以下为本月各品类目标达成与环比变动情况，数字即真相，不粉饰、不隐瞒。",
                        italic=True, color=COLORS["text_muted"], indent=False)

        # 构建表头
        cols = len(self.data.categories) + 1
        table = self.doc.add_table(rows=5, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        headers = ["经营指标"] + [c.name for c in self.data.categories]
        self._make_table_header(table, headers)

        # 准备数据行
        row_target = ["本月目标"] + [f"{c.target:,.0f}{c.unit}" for c in self.data.categories]
        row_actual = ["实际达成"] + [f"{c.actual:,.0f}{c.unit}" for c in self.data.categories]
        row_rate = ["达成率"]
        row_mom = ["环比变动"]

        rate_aligns = [WD_ALIGN_PARAGRAPH.CENTER] * cols
        rate_shadings = [None]
        rate_bolds = [False]

        for c in self.data.categories:
            rate = c.actual / c.target if c.target else 0
            rate_str = f"{rate*100:.2f}%"
            row_rate.append(rate_str)
            rate_shadings.append(self._rate_bg(rate))
            rate_bolds.append(True)

            if c.last_month == 0:
                mom_str = "-" if c.actual == 0 else "新增长"
            else:
                mom = (c.actual - c.last_month) / c.last_month
                sign = "+" if mom >= 0 else ""
                mom_str = f"{sign}{mom*100:.1f}%"
            row_mom.append(mom_str)
        
        # 保存达成率和环比数据供后续着色使用
        self._rates = [c.actual / c.target if c.target else 0 for c in self.data.categories]
        self._moms = []
        for c in self.data.categories:
            if c.last_month == 0:
                self._moms.append(None)
            else:
                self._moms.append((c.actual - c.last_month) / c.last_month)

        # 写入表格
        for ridx, vals in enumerate([row_target, row_actual, row_rate, row_mom], 1):
            row = table.rows[ridx]
            for j, text in enumerate(vals):
                cell = row.cells[j]
                set_cell_border(cell, color=COLORS["border"], size="4")
                set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
                if j == 0:
                    set_cell_shading(cell, COLORS["bg_light"])

                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                is_bold = (ridx == 3)  # 达成率行加粗
                c = COLORS["text"]
                if ridx == 4 and j > 0:
                    # 环比变动颜色（精确判断，避免"-"被误判）
                    mom_val = self._moms[j-1] if (j-1) < len(self._moms) else None
                    if mom_val is not None:
                        c = COLORS["success"] if mom_val >= 0 else COLORS["danger"]
                    else:
                        c = COLORS["text"]
                set_run_font(run, size=9.5, bold=is_bold, color=c)

        # 达成率行单独加背景色
        for j, shd in enumerate(rate_shadings):
            if shd:
                set_cell_shading(table.rows[3].cells[j], shd)

        # 调整列宽
        for row in table.rows:
            row.cells[0].width = Cm(2.5)
            for j in range(1, cols):
                row.cells[j].width = Cm(2.0)

        self.doc.add_paragraph()

        # 亮点与短板文字分析
        if self.data.highlights:
            self._add_heading("2.1 亮点品类", level=2, color=COLORS["success"])
            for h in self.data.highlights:
                self._add_bullet(h, color=COLORS["success"], icon="▲")

        if self.data.weaknesses:
            self._add_heading("2.2 短板品类", level=2, color=COLORS["danger"])
            for w in self.data.weaknesses:
                self._add_bullet(w, color=COLORS["danger"], icon="▼")

        self.doc.add_paragraph()

    # ==================== 第三章：费用与利润分析 ====================
    def add_cost_profit(self):
        self._add_heading("三、费用与利润分析", level=1)
        self._add_para("以「肌肉型经营」审视每一笔费用，去除赘肉，保留创造价值的核心投入。",
                        italic=True, color=COLORS["text_muted"], indent=False)

        # 费用明细表
        fee_table = self.doc.add_table(rows=len(self.data.costs) + 2, cols=4)
        fee_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._make_table_header(fee_table, ["费用大类", "明细项目", "金额（元）", "费用审视"],
                                color=COLORS["secondary"])

        for i, item in enumerate(self.data.costs, 1):
            aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                      WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
            shadings = [COLORS["bg_light"], None, None, None]
            bolds = [True, False, True, False]
            values = [item.category, item.name, f"{item.amount:,.0f}", item.review]
            self._make_table_row(fee_table, i, values, aligns, shadings, bolds)

        # 合计行
        total_idx = len(self.data.costs) + 1
        total_row = fee_table.rows[total_idx]
        for j in range(4):
            cell = total_row.cells[j]
            set_cell_shading(cell, "FFF9C4")
            set_cell_border(cell, color=COLORS["accent"], size="6")
            set_cell_margins(cell, top=100, bottom=100)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j < 3 else WD_ALIGN_PARAGRAPH.LEFT
            fee_rate = self.data.cost_total / self.data.gmv * 100 if self.data.gmv else 0
            rebate_ratio = self.data.cost_total / max(self.data.rebate_amount, 0.01) * 100 if self.data.rebate_amount else 0
            texts = ["", "费用合计", f"¥{self.data.cost_total:,.0f}",
                         f"费用率 {fee_rate:.1f}% | 占返利 {rebate_ratio:.1f}%"]
            run = p.add_run(texts[j])
            set_run_font(run, size=10.5, bold=True, color=COLORS["danger"] if self.data.cost_total > self.data.rebate_amount else COLORS["primary"])

        for row in fee_table.rows:
            row.cells[0].width = Cm(2.2)
            row.cells[1].width = Cm(3.5)
            row.cells[2].width = Cm(2.5)
            row.cells[3].width = Cm(6.5)

        self.doc.add_paragraph()

        # 净利润核算表
        self._add_heading("3.1 净利润核算", level=2)
        profit_table = self.doc.add_table(rows=4, cols=2)
        profit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._make_table_header(profit_table, ["核算项目", "金额（元）"], color=COLORS["accent"])

        items = [
            ("① 返利收入（销售额 × 返利率）", self.data.rebate_amount),
            ("② 费用合计", self.data.cost_total),
            ("③ 净利润 = ① - ②", self.data.profit),
        ]
        for i, (label, val) in enumerate(items, 1):
            cell0 = profit_table.rows[i].cells[0]
            cell1 = profit_table.rows[i].cells[1]
            set_cell_border(cell0, color=COLORS["border"], size="4")
            set_cell_border(cell1, color=COLORS["border"], size="4")
            set_cell_margins(cell0, top=120, bottom=120)
            set_cell_margins(cell1, top=120, bottom=120)

            p0 = cell0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r0 = p0.add_run(label)
            is_bold = (i == 3)
            c = COLORS["danger"] if i == 3 and self.data.profit < 0 else COLORS["text"]
            set_run_font(r0, size=11, bold=is_bold, color=c)

            set_cell_shading(cell1, "F3E5F5" if i == 3 else COLORS["bg_light"])
            p1 = cell1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p1.add_run(f"¥{val:,.0f}")
            set_run_font(r1, size=12, bold=is_bold, color=c)

        for row in profit_table.rows:
            row.cells[0].width = Cm(9)
            row.cells[1].width = Cm(5)

        self.doc.add_paragraph()

    # ==================== 第四章：问题拆解 ====================
    def add_problems(self):
        self._add_heading("四、问题拆解与真因探究", level=1)
        self._add_para("不将原因归咎于外部，而是反复追问「为什么」，直到找到可改善的真因。",
                        italic=True, color=COLORS["text_muted"], indent=False)

        for prob in self.data.problems:
            self._add_bullet(prob, icon="●")

        if self.data.five_whys:
            self.doc.add_paragraph()
            self._add_heading("4.1 五个为什么", level=2)
            why_table = self.doc.add_table(rows=len(self.data.five_whys), cols=2)
            why_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for i, (q, a) in enumerate(self.data.five_whys):
                row = why_table.rows[i]
                cell0 = row.cells[0]
                cell1 = row.cells[1]
                set_cell_shading(cell0, "FFE4E6")
                set_cell_border(cell0, color=COLORS["border"], size="4")
                set_cell_border(cell1, color=COLORS["border"], size="4")
                set_cell_margins(cell0, top=100, bottom=100, left=120, right=120)
                set_cell_margins(cell1, top=100, bottom=100, left=120, right=120)

                p0 = cell0.paragraphs[0]
                p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r0 = p0.add_run(q)
                set_run_font(r0, size=10, bold=True, color=COLORS["danger"])

                p1 = cell1.paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r1 = p1.add_run(a)
                set_run_font(r1, size=10, color=COLORS["text"])

            for row in why_table.rows:
                row.cells[0].width = Cm(5)
                row.cells[1].width = Cm(9)

        self.doc.add_paragraph()

    # ==================== 第五章：改善行动计划 ====================
    def add_actions(self):
        self._add_heading("五、改善行动计划", level=1)
        self._add_para("目标既定，无论遇到什么困难都绝不退缩，以可量化、可检查的标准立即执行。",
                        italic=True, color=COLORS["text_muted"], indent=False)

        # 按分类分组
        categories = {}
        for act in self.data.actions:
            categories.setdefault(act.category, []).append(act)

        for cat_name, acts in categories.items():
            self._add_heading(cat_name, level=2)
            act_table = self.doc.add_table(rows=len(acts) + 1, cols=4)
            act_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            self._make_table_header(act_table, ["序号", "改善动作", "量化标准（可检查）", "责任人/周期"],
                                    color=COLORS["accent"])

            for i, act in enumerate(acts, 1):
                values = [str(i), act.action, act.standard, act.owner]
                aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                          WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER]
                self._make_table_row(act_table, i, values, aligns)

            for row in act_table.rows:
                row.cells[0].width = Cm(1.2)
                row.cells[1].width = Cm(4.5)
                row.cells[2].width = Cm(5.5)
                row.cells[3].width = Cm(3)

            self.doc.add_paragraph()

    # ==================== 第六章：运营数据表现 ====================
    def add_operation(self):
        self._add_heading("六、运营数据表现", level=1)

        op_table = self.doc.add_table(rows=len(self.data.operation_metrics) + 1, cols=4)
        op_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._make_table_header(op_table, ["运营指标", "上月数据", "本月数据", "诊断说明"],
                                color=COLORS["secondary"])

        for i, m in enumerate(self.data.operation_metrics, 1):
            values = [m.name, m.last_month, m.current_month, m.diagnosis]
            aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                      WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
            shadings = [COLORS["bg_light"], None, None, None]
            bolds = [True, False, False, False]
            self._make_table_row(op_table, i, values, aligns, shadings, bolds)

        for row in op_table.rows:
            row.cells[0].width = Cm(3)
            row.cells[1].width = Cm(2.5)
            row.cells[2].width = Cm(2.5)
            row.cells[3].width = Cm(6.5)

        self.doc.add_paragraph()

    # ==================== 第七章：下月经营计划 ====================
    def add_next_month(self):
        self._add_heading("七、下月经营计划", level=1)
        self._add_para("胸中怀有强烈的愿望，设立具体目标，让每位伙伴都清楚自己的方向。",
                        italic=True, color=COLORS["text_muted"], indent=False)

        goal_table = self.doc.add_table(rows=len(self.data.next_month_goals) + 1, cols=4)
        goal_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._make_table_header(goal_table, ["目标维度", "基础目标", "挑战目标", "达成策略简述"],
                                color=COLORS["primary"])

        for i, g in enumerate(self.data.next_month_goals, 1):
            values = [g["维度"], g["基础目标"], g["挑战目标"], g["策略"]]
            aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                      WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
            shadings = [COLORS["bg_accent"], "E3F2FD", "E3F2FD", None]
            self._make_table_row(goal_table, i, values, aligns, shadings)

        for row in goal_table.rows:
            row.cells[0].width = Cm(3)
            row.cells[1].width = Cm(3)
            row.cells[2].width = Cm(3)
            row.cells[3].width = Cm(6)

        self.doc.add_paragraph()

    # ==================== 第八章：经营者日记 ====================
    def add_reflection(self):
        self._add_heading("八、经营者心灵日记", level=1)
        self._add_para("诚实面对经营结果，感谢团队付出，坚定下月决心。",
                        italic=True, color=COLORS["text_muted"], indent=False)

        # 用带底色边框的段落模拟日记卡片
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_shading(cell, COLORS["bg_warm"])
        set_cell_border(cell, color=COLORS["accent"], size="6")
        set_cell_margins(cell, top=200, bottom=200, left=200, right=200)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.8
        run = p.add_run(self.data.reflection)
        set_run_font(run, size=11, color=COLORS["secondary"])

        self.doc.add_paragraph()

    # ==================== 结尾 ====================
    def add_footer(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(30)
        run = p.add_run("—— 敬天爱人 · 追求卓越 ——")
        set_run_font(run, size=13, italic=True, color=COLORS["accent"])

    # ==================== 组装生成 ====================
    def generate(self, output_path: str):
        self.add_cover()
        self.add_summary()
        self.add_sales_analysis()
        self.add_cost_profit()
        self.add_problems()
        self.add_actions()
        self.add_operation()
        self.add_next_month()
        self.add_reflection()
        self.add_footer()
        self.doc.save(output_path)
        print(f"✅ 经营复盘报告已生成：{output_path}")


# ==================== 主入口（示例） ====================
if __name__ == "__main__":
    # 从示例数据加载（实际使用时替换为用户自己的数据）
    from examples.sample_data import get_sample_data
    data = get_sample_data()
    report = StoreReviewReport(data)
    report.generate("./门店经营复盘报告_示例.docx")
