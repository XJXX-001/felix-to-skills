#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
苹果门店月度经营复盘模板（填空版）
融合稻盛和夫经营哲学
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'single')
        edge_el.set(qn('w:sz'), '4')
        edge_el.set(qn('w:color'), '000000')
        tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin = OxmlElement(f'w:{edge}')
        margin.set(qn('w:w'), str(val))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)

def add_heading_custom(doc, text, level=1, color="1F4E79"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.bold = True
    run.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
    sizes = {1: 18, 2: 15, 3: 13}
    run.font.size = Pt(sizes.get(level, 12))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_para(doc, text, bold=False, italic=False, size=10.5, color="333333", align=WD_ALIGN_PARAGRAPH.LEFT, indent=True, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_fill_line(doc, label, lines=1, hint=""):
    """添加带标签的填写区域（用浅色边框表格替代下划线，更美观统一）"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    
    run = p.add_run(label)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    if hint:
        run2 = p.add_run(f"  ({hint})")
        run2.font.name = 'Microsoft YaHei'
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run2.font.size = Pt(9)
        run2.font.italic = True
        run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    for _ in range(lines):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_shading(cell, "FAFAFA")
        set_cell_border(cell)
        set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
        
        p_cell = cell.paragraphs[0]
        p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 留白，仅一个提示性空 run
        run_cell = p_cell.add_run("")
        run_cell.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        
        for row in table.rows:
            row.cells[0].width = Cm(14)

def add_philosophy_box(doc, title, content, icon="📖"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF8E7")
    set_cell_border(cell)
    set_cell_margins(cell)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{icon} 【稻盛和夫·{title}】")
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
    
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = p2.add_run(content)
    run2.font.name = 'Microsoft YaHei'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x5D, 0x40, 0x30)
    p2.paragraph_format.line_spacing = 1.5
    
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()

def create_template():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(10.5)
    
    # ========== 封面 ==========
    doc.add_paragraph()
    doc.add_paragraph()
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("苹果门店月度经营复盘报告")
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— 融合稻盛和夫经营哲学 · 店长填写版 ——")
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 基本信息填写区
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_items = [
        ("门店名称", "如：示例店A / 示例店B"),
        ("店长姓名", "请填写"),
        ("复盘月份", "如：2026年4月"),
        ("填报日期", "如：2026年5月5日"),
        ("门店人数", "如：3人（含店长）"),
    ]
    for i, (label, placeholder) in enumerate(info_items):
        row = info_table.rows[i]
        cell0 = row.cells[0]
        cell1 = row.cells[1]
        set_cell_shading(cell0, "E7F0F7")
        set_cell_border(cell0)
        set_cell_border(cell1)
        
        p0 = cell0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.name = 'Microsoft YaHei'
        r0._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        r0.font.size = Pt(11)
        r0.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        
        p1 = cell1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(placeholder)
        r1.font.name = 'Microsoft YaHei'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        
        for cell in row.cells:
            set_cell_margins(cell, top=100, bottom=100)
    
    for row in info_table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(8)
    
    doc.add_page_break()
    
    # ========== 一、门店经营结果与成本·净利呈现 ==========
    add_heading_custom(doc, "一、门店经营结果与成本·净利呈现", 1)
    
    add_para(doc, "请在下表中如实填写本月经营结果、成本结构与净利润。数字即真相，不粉饰、不隐瞒。",
        italic=True, color="666666", indent=False)
    doc.add_paragraph()
    
    # 表1：经营结果（收入侧）
    tbl1 = doc.add_table(rows=7, cols=4)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers1 = ["经营指标", "本月目标", "实际达成", "达成率 / 差异分析"]
    for i, h in enumerate(headers1):
        cell = tbl1.rows[0].cells[i]
        set_cell_shading(cell, "1F4E79")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10.5)
    
    rows1_data = [
        ("总销售毛利（元）", "", ""),
        ("手机销量（台）", "", ""),
        ("ACS（单）", "", ""),
        ("二手机回收毛利（元）", "", ""),
        ("会员新增（人）", "", ""),
        ("其他收入（元）", "0", "如有"),
    ]
    for i, (metric, target_hint, actual_hint) in enumerate(rows1_data, 1):
        row = tbl1.rows[i]
        for j in range(4):
            cell = row.cells[j]
            set_cell_border(cell)
            if j == 0:
                set_cell_shading(cell, "E7F0F7")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(metric)
                run.font.bold = True
            else:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                hints = [target_hint, actual_hint, "请填写"]
                hint_text = hints[j-1] if hints[j-1] else ""
                run = p.add_run(hint_text)
                run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(10)
            set_cell_margins(cell, top=100, bottom=100)
    
    for row in tbl1.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(3)
        row.cells[2].width = Cm(3)
        row.cells[3].width = Cm(5.5)
    
    doc.add_paragraph()
    
    # 表2：成本结构（费用侧）
    tbl2 = doc.add_table(rows=9, cols=3)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers2 = ["费用项目", "本月金额（元）", "费用审视（店长填写）"]
    for i, h in enumerate(headers2):
        cell = tbl2.rows[0].cells[i]
        set_cell_shading(cell, "2E7D32")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10.5)
    
    rows2_data = [
        ("房租", "刚性支出，关注坪效"),
        ("水电费", "节能空间自查"),
        ("人员薪资（含社保）", "人均产出是否匹配"),
        ("装修/折旧摊销", "一次性投入分摊"),
        ("收款手续费", "费率优化空间"),
        ("物料/包装/耗材", "小额费用亦不可忽视"),
        ("物流/其他杂费", "逐项核对必要性"),
    ]
    for i, (item, review) in enumerate(rows2_data, 1):
        row = tbl2.rows[i]
        for j in range(3):
            cell = row.cells[j]
            set_cell_border(cell)
            if j == 0:
                set_cell_shading(cell, "E8F5E9")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(item)
                run.font.bold = True
            elif j == 1:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("")
                run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            else:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(review)
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                run.font.italic = True
                run.font.size = Pt(9.5)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(10)
            set_cell_margins(cell, top=90, bottom=90)
    
    # 合计行
    total_row = tbl2.add_row()
    for j in range(3):
        cell = total_row.cells[j]
        set_cell_shading(cell, "FFF9C4")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j < 2 else WD_ALIGN_PARAGRAPH.LEFT
        texts = ["", "费用合计", "费用率 = 费用合计 ÷ 总销售毛利"]
        run = p.add_run(texts[j])
        run.font.bold = True
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if j == 2:
            run.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
            run.font.size = Pt(9.5)
    
    for row in tbl2.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(6.5)
    
    doc.add_paragraph()
    
    # 表3：净利润呈现
    tbl3 = doc.add_table(rows=4, cols=2)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers3 = ["核算项目", "金额（元）"]
    for i, h in enumerate(headers3):
        cell = tbl3.rows[0].cells[i]
        set_cell_shading(cell, "5E35B1")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(11)
    
    rows3_data = [
        ("① 总销售毛利（来自上方表1）", ""),
        ("② 费用合计（来自上方表2）", ""),
        ("③ 净利润 = ① - ②", ""),
    ]
    for i, (item, val) in enumerate(rows3_data, 1):
        row = tbl3.rows[i]
        for j in range(2):
            cell = row.cells[j]
            set_cell_border(cell)
            if j == 0:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(item)
                run.font.bold = True
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                if i == 3:
                    run.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
                    run.font.size = Pt(11)
            else:
                set_cell_shading(cell, "F3E5F5")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("")
                run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                if i == 3:
                    run.font.bold = True
                    run.font.size = Pt(12)
            set_cell_margins(cell, top=120, bottom=120)
    
    for row in tbl3.rows:
        row.cells[0].width = Cm(9)
        row.cells[1].width = Cm(5)
    
    doc.add_paragraph()
    
    # 提示段落
    p_hint = doc.add_paragraph()
    p_hint.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_hint.paragraph_format.line_spacing = 1.5
    p_hint.paragraph_format.space_after = Pt(6)
    run = p_hint.add_run("💡 提示：净利润是门店经营的「终极成绩单」。若净利润为负，请重点关注：")
    run.font.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    bullets = [
        "• 费用率是否超过警戒线（建议控制在总毛利的60%以内）；",
        "• 各品类达成率是否均衡，是否存在「高毛利品类」短板；",
        "• 是否有可削减的非必要费用（肌肉型经营原则）。",
    ]
    for b in bullets:
        p_b = doc.add_paragraph()
        p_b.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_b.paragraph_format.line_spacing = 1.5
        p_b.paragraph_format.space_after = Pt(2)
        p_b.paragraph_format.left_indent = Cm(0.8)
        run = p_b.add_run(b)
        run.font.size = Pt(10)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # ========== 二、经营目标与实际达成 ==========
    add_heading_custom(doc, "二、经营目标与实际达成（核心数据表）", 1)
    
    add_para(doc, "请在下表中填入本月目标（公司下达）与实际达成数据，并计算达成率：", 
        italic=True, color="666666", indent=False)
    doc.add_paragraph()
    
    # 核心数据对比表
    table = doc.add_table(rows=9, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["经营指标", "本月目标", "实际达成", "达成率", "差异分析（店长填写）"]
    row0 = table.rows[0]
    for i, h in enumerate(headers):
        cell = row0.cells[i]
        set_cell_shading(cell, "1F4E79")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10.5)
    
    metrics = [
        "手机销量（台）",
        "ACS（AppleCare服务）",
        "二手机回收毛利（元）",
        "总销售毛利（元）",
        "回收台数（台）",
        "会员新增（人）",
        "16-128清库（台）",
        "MacBook Air M4清库（台）",
    ]
    
    for i, metric in enumerate(metrics, 1):
        row = table.rows[i]
        for j in range(5):
            cell = row.cells[j]
            set_cell_border(cell)
            if j == 0:
                set_cell_shading(cell, "E7F0F7")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(metric)
                run.font.bold = True
            else:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                placeholder = ["", "", "", "请填写差异原因及反思"][j-1]
                run = p.add_run(placeholder)
                run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
                if j == 3:
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(10)
        for cell in row.cells:
            set_cell_margins(cell, top=120, bottom=120)
    
    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(3)
        row.cells[2].width = Cm(3)
        row.cells[3].width = Cm(2.5)
        row.cells[4].width = Cm(5)
    
    doc.add_paragraph()
    
    # ========== 三、数据洞察 ==========
    add_heading_custom(doc, "三、数据洞察与品类分析（数字经营）", 1)
    
    add_philosophy_box(doc, "京瓷会计七原则·完美主义", 
        "数据不会说谎。请如实面对每一个数字，不粉饰、不推诿，"
        "「完美无缺」地呈现经营真相。", "🔍")
    
    add_heading_custom(doc, "3.1 亮点品类（发扬光大）", 2)
    add_para(doc, "请列出本月达成最好、最值得骄傲的1-3个指标，并分析成功经验（可复制的方法）：", 
        italic=True, color="666666", indent=False)
    add_fill_line(doc, "① 亮点指标：", lines=2, hint="如：MacBook Air M4清库达成150%，原因是...")
    add_fill_line(doc, "② 亮点指标：", lines=2)
    add_fill_line(doc, "③ 亮点指标：", lines=2)
    
    add_heading_custom(doc, "3.2 短板品类（直面问题）", 2)
    add_para(doc, "请列出本月未达标、下滑最严重的1-3个指标，不回避、不找借口：", 
        italic=True, color="666666", indent=False)
    add_fill_line(doc, "① 短板指标：", lines=2, hint="如：手机销量达成仅70%，未达标原因...")
    add_fill_line(doc, "② 短板指标：", lines=2)
    add_fill_line(doc, "③ 短板指标：", lines=2)
    
    add_heading_custom(doc, "3.3 关键运营数据", 2)
    add_para(doc, "请填写以下门店过程数据，用于分析客流与转化：", italic=True, color="666666", indent=False)
    
    op_table = doc.add_table(rows=5, cols=4)
    op_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    op_headers = ["运营指标", "上月数据", "本月数据", "环比变化分析"]
    row0 = op_table.rows[0]
    for i, h in enumerate(op_headers):
        cell = row0.cells[i]
        set_cell_shading(cell, "2E7D32")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10.5)
    
    op_metrics = ["客流量（人）", "成交转化率（%）", "客单价（元）", "连带率（件/单）"]
    for i, metric in enumerate(op_metrics, 1):
        row = op_table.rows[i]
        for j in range(4):
            cell = row.cells[j]
            set_cell_border(cell)
            if j == 0:
                set_cell_shading(cell, "E8F5E9")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(metric)
                run.font.bold = True
            else:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("")
                run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(10)
        for cell in row.cells:
            set_cell_margins(cell, top=100, bottom=100)
    
    for row in op_table.rows:
        for cell in row.cells:
            cell.width = Cm(3.75)
    
    doc.add_paragraph()
    
    # ========== 四、问题真因探究 ==========
    add_heading_custom(doc, "四、问题真因探究（5个为什么·追究本质）", 1)
    
    add_philosophy_box(doc, "经营十二条·第4条", 
        "付出不亚于任何人的努力——成功没有捷径，唯有深入现场、反复琢磨，"
        "才能找到真正的突破口。", "🔥")
    
    add_para(doc, "请针对最核心的1个未达标指标，运用「5个为什么」连续追问，找到真因：", 
        bold=True, color="1F4E79", indent=False)
    
    why_table = doc.add_table(rows=6, cols=2)
    why_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    why_labels = [
        "第1问：问题现象是什么？",
        "第2问：为什么会发生？",
        "第3问：为什么会这样？",
        "第4问：更深一层的原因是什么？",
        "第5问：根本真因是什么？",
        "真因归类（可多选）"
    ]
    why_hints = [
        "描述具体数据和事实",
        "表面原因",
        "中间层原因",
        "接近本质的原因",
        "可执行、可改善的根本原因",
        "□ 人员能力  □ 执行力  □ 流程制度  □ 资源支持  □ 外部环境"
    ]
    
    for i in range(6):
        row = why_table.rows[i]
        cell0 = row.cells[0]
        cell1 = row.cells[1]
        set_cell_border(cell0)
        set_cell_border(cell1)
        set_cell_shading(cell0, "FFEBEE")
        set_cell_margins(cell0)
        set_cell_margins(cell1)
        
        p0 = cell0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r0 = p0.add_run(why_labels[i])
        r0.font.bold = True
        r0.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
        r0.font.name = 'Microsoft YaHei'
        r0._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        r0.font.size = Pt(10.5)
        
        p1 = cell1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(why_hints[i])
        r1.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB) if i < 5 else RGBColor(0x33, 0x33, 0x33)
        r1.font.name = 'Microsoft YaHei'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        r1.font.size = Pt(10)
        if i == 5:
            r1.font.bold = True
    
    for row in why_table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(9.5)
    
    doc.add_paragraph()
    
    # ========== 五、改善行动计划 ==========
    add_heading_custom(doc, "五、改善行动计划（PDCA·立即执行）", 1)
    
    add_philosophy_box(doc, "经营十二条·第7条·第10条", 
        "经营取决于坚强的意志——目标既定，无论遇到什么困难，绝不退缩。\n"
        "不断从事创造性的工作——明天胜过今天，后天胜过明天。", "💪")
    
    add_para(doc, "请针对上述真因，制定至少3条可量化、可检查、有期限的改善动作：", 
        bold=True, color="1F4E79", indent=False)
    doc.add_paragraph()
    
    action_table = doc.add_table(rows=4, cols=5)
    action_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    a_headers = ["序号", "改善动作", "量化标准（可检查）", "责任人", "完成期限"]
    row0 = action_table.rows[0]
    for i, h in enumerate(a_headers):
        cell = row0.cells[i]
        set_cell_shading(cell, "C62828")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10.5)
    
    for i in range(1, 4):
        row = action_table.rows[i]
        for j in range(5):
            cell = row.cells[j]
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            text = str(i) if j == 0 else ""
            run = p.add_run(text)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
        for cell in row.cells:
            set_cell_margins(cell, top=120, bottom=120)
    
    for row in action_table.rows:
        row.cells[0].width = Cm(1.2)
        row.cells[1].width = Cm(4.5)
        row.cells[2].width = Cm(4.5)
        row.cells[3].width = Cm(2.5)
        row.cells[4].width = Cm(2.5)
    
    doc.add_paragraph()
    
    add_heading_custom(doc, "5.1 过程管理承诺", 2)
    add_para(doc, "请勾选并承诺本月过程管理动作：", bold=True, indent=False)
    
    checks = [
        "□ 每日晨会：明确当日手机/ACS/回收/会员目标，分解到人",
        "□ 每日晚会：核对当日达成，差额次日补足，登记《日清表》",
        "□ 每周复盘：周五晚召开周复盘会，逐项核对目标完成情况",
        "□ 重点产品：MacBook Air M4、16-128清库每日主推，100%讲解",
        "□ 会员运营：每日新增会员目标分解，到店顾客100%扫码入会",
        "□ 回收业务：每台二手机回收必登记，毛利透明化",
        "□ 店长巡检：每日巡检3次，纠偏连带销售、ACS推荐遗漏",
    ]
    for c in checks:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.8
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(c)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # ========== 六、下月经营计划 ==========
    add_heading_custom(doc, "六、下月经营计划（胸中怀有强烈的愿望）", 1)
    
    add_philosophy_box(doc, "经营十二条·第3条", 
        "胸中怀有强烈的愿望——要怀有渗透到潜意识的强烈而持久的愿望，"
        "「无论如何也要达成」的执着意念，是经营成功的第一推动力。", "⭐")
    
    add_para(doc, "请填写下月核心目标（由公司下达或自主挑战）：", bold=True, indent=False)
    doc.add_paragraph()
    
    goal_table = doc.add_table(rows=9, cols=4)
    goal_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    g_headers = ["经营指标", "基础目标", "挑战目标", "达成策略简述"]
    row0 = goal_table.rows[0]
    for i, h in enumerate(g_headers):
        cell = row0.cells[i]
        set_cell_shading(cell, "1565C0")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10.5)
    
    goal_metrics = [
        "手机销量（台）",
        "ACS（单）",
        "二手机回收毛利（元）",
        "总销售毛利（元）",
        "回收台数（台）",
        "会员新增（人）",
        "16-128清库（台）",
        "MacBook Air M4清库（台）",
    ]
    
    for i, metric in enumerate(goal_metrics, 1):
        row = goal_table.rows[i]
        for j in range(4):
            cell = row.cells[j]
            set_cell_border(cell)
            if j == 0:
                set_cell_shading(cell, "E3F2FD")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(metric)
                run.font.bold = True
            else:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("")
                run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(10)
        for cell in row.cells:
            set_cell_margins(cell, top=100, bottom=100)
    
    for row in goal_table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(3)
        row.cells[2].width = Cm(3)
        row.cells[3].width = Cm(5)
    
    doc.add_paragraph()
    add_fill_line(doc, "下月重点突破方向（1-2句话概括）：", lines=2, hint="如：聚焦手机销量回升+MacBook清库冲刺")
    
    # ========== 七、经营者心灵日记 ==========
    add_heading_custom(doc, "七、经营者心灵日记（敬天爱人·保持乐观向上）", 1)
    
    add_philosophy_box(doc, "经营十二条·第11条·第12条", 
        "以关怀之心、诚实处事——买卖是双方的，生意各方都得利，皆大欢喜。\n"
        "保持乐观向上的态度——抱着梦想与希望，以坦诚之心处世。", "❤️")
    
    add_para(doc, "请店长用3-5句话，诚实记录本月的经营感悟、对团队的感谢、以及对下月的决心：", 
        italic=True, color="666666", indent=False)
    
    # 用带边框的表格卡片替代下划线，更美观且便于填写
    diary_table = doc.add_table(rows=1, cols=1)
    diary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = diary_table.cell(0, 0)
    set_cell_shading(cell, "FFFBEB")
    set_cell_border(cell)
    set_cell_margins(cell, top=200, bottom=200, left=200, right=200)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 2.0
    run_hint = p.add_run("（请在此区域手写或输入经营者日记内容）")
    run_hint.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    run_hint.font.name = 'Microsoft YaHei'
    run_hint._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run_hint.font.size = Pt(10)
    run_hint.font.italic = True
    
    # ========== 附录：各门店目标参考 ==========
    doc.add_page_break()
    add_heading_custom(doc, "附录：各门店月度销售目标参考（公司下达）", 1)
    
    add_para(doc, "以下为本月度各门店目标值，供店长对照参考。请找到自己所在门店，将目标填入第二部分的表格中。", 
        italic=True, color="666666", indent=False)
    doc.add_paragraph()
    
    ref_table = doc.add_table(rows=8, cols=10)
    ref_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ref_headers = ["门店", "店长", "手机", "ACS", "回收毛利", "总毛利", "回收台数", "会员", "16清库", "Mac清库"]
    row0 = ref_table.rows[0]
    for i, h in enumerate(ref_headers):
        cell = row0.cells[i]
        set_cell_shading(cell, "5E35B1")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(9.5)
    
    ref_data = [
        ("示例店A", "张三", "80", "12", "15,000", "60,000", "30", "40", "5", "5"),
        ("示例店B", "李四", "60", "8", "10,000", "50,000", "20", "35", "5", "5"),
        ("示例店C", "王五", "200", "30", "35,000", "200,000", "70", "110", "10", "10"),
        ("示例店D", "赵六", "140", "22", "25,000", "140,000", "55", "80", "10", "10"),
        ("示例店E", "孙七", "100", "14", "12,000", "90,000", "35", "60", "8", "5"),
        ("示例店F", "周八", "70", "10", "9,000", "65,000", "28", "40", "6", "5"),
        ("合计", "-", "650", "96", "106,000", "605,000", "238", "365", "44", "40"),
    ]
    
    for i, row_data in enumerate(ref_data, 1):
        row = ref_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            set_cell_border(cell)
            if i == 7:
                set_cell_shading(cell, "FFF9C4")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(9)
            if i == 7:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
        for cell in row.cells:
            set_cell_margins(cell, top=60, bottom=60)
    
    # ========== 附录：复盘工具说明 ==========
    doc.add_paragraph()
    add_heading_custom(doc, "附录：复盘填写说明", 2)
    
    tips = [
        "1. 本模板融合稻盛和夫「经营十二条」与「京瓷会计七原则」，请店长认真、诚实填写；",
        "2. 「5个为什么」部分请务必追根溯源，禁止将原因简单归咎于「客人少」「市场差」等外部因素；",
        "3. 改善行动计划必须满足 SMART 原则：具体的(S)、可衡量的(M)、可达成的(A)、相关的(R)、有时限的(T)；",
        "4. 每月5日前提交上月复盘报告至区域经理，作为月度绩效评定依据；",
        "5. 复盘不是为了追责，而是为了共同成长。保持乐观、坦诚面对，才是真正的经营者姿态。",
    ]
    for tip in tips:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.8
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(tip)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    # 结尾
    doc.add_paragraph()
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_p.add_run("—— 敬天爱人 · 追求卓越 ——")
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
    end_p.paragraph_format.space_before = Pt(20)
    
    output_path = "./苹果门店月度复盘模板_店长填写版.docx"
    doc.save(output_path)
    print(f"✅ Word模板已生成：{output_path}")

if __name__ == "__main__":
    create_template()
